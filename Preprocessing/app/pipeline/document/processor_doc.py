"""
Document Preprocessing Branch
--------------------------------
Steps  (matches diagram Image 1):
  1. Text extract  – PyMuPDF (primary) with pdfplumber fallback
  2. Layout parse  – headings, tables, section structure
  3. Doc metadata  – author, edit history, creation date, software
"""

from __future__ import annotations
import hashlib
import io
import logging
import re
from pathlib import Path
from typing import Union

from Preprocessing.app.models.feature_object import (
    DocMetadata,
    DocumentData,
    InputType,
    LayoutInfo,
    NormalizedFeatureObject,
    TextExtract,
)

logger = logging.getLogger(__name__)


# ══════════════════════════════════════════════════════
# 1.  Text extraction
# ══════════════════════════════════════════════════════

def extract_text(file_bytes: bytes, filename: str) -> TextExtract:
    """
    Extract visible text from a document.
    Tries PyMuPDF first; falls back to pdfplumber for PDFs.
    """
    ext = Path(filename).suffix.lower()

    # ── PDF path ──
    if ext == ".pdf":
        return _extract_pdf(file_bytes)

    # ── DOCX / DOC path ──
    if ext in (".docx", ".doc"):
        return _extract_docx(file_bytes)

    # ── Plain text ──
    if ext in (".txt", ".md", ".rst"):
        try:
            text = file_bytes.decode("utf-8", errors="replace")
            return TextExtract(raw_text=text, page_count=1, extraction_method="plain")
        except Exception as e:
            logger.warning(f"Plain text decode failed: {e}")
            return TextExtract(raw_text="", extraction_method="plain")

    return TextExtract(raw_text="", extraction_method="unsupported")


def _extract_pdf(file_bytes: bytes) -> TextExtract:
    """Try PyMuPDF; fall back to pdfplumber."""
    # Primary: PyMuPDF
    try:
        import fitz  # PyMuPDF

        doc = fitz.open(stream=file_bytes, filetype="pdf")
        pages: list[str] = []
        for page in doc:
            pages.append(page.get_text())
        doc.close()

        raw_text = "\n".join(pages)
        return TextExtract(
            raw_text=raw_text,
            page_count=len(pages),
            extraction_method="pymupdf",
        )
    except ImportError:
        logger.info("PyMuPDF not installed, trying pdfplumber")
    except Exception as e:
        logger.warning(f"PyMuPDF extraction failed: {e}")

    # Fallback: pdfplumber
    try:
        import pdfplumber

        pages = []
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            for page in pdf.pages:
                text = page.extract_text() or ""
                pages.append(text)

        raw_text = "\n".join(pages)
        return TextExtract(
            raw_text=raw_text,
            page_count=len(pages),
            extraction_method="pdfplumber",
        )
    except ImportError:
        logger.warning("pdfplumber not installed")
    except Exception as e:
        logger.warning(f"pdfplumber extraction failed: {e}")

    return TextExtract(raw_text="", extraction_method="failed")


def _extract_docx(file_bytes: bytes) -> TextExtract:
    """Extract text from DOCX using python-docx."""
    try:
        import docx

        doc = docx.Document(io.BytesIO(file_bytes))
        paragraphs = [p.text for p in doc.paragraphs]
        raw_text = "\n".join(paragraphs)
        # Word documents don't have "pages" — approximate by word count
        word_count = len(raw_text.split())
        approx_pages = max(1, word_count // 300)
        return TextExtract(
            raw_text=raw_text,
            page_count=approx_pages,
            extraction_method="python-docx",
        )
    except ImportError:
        logger.warning("python-docx not installed")
    except Exception as e:
        logger.warning(f"DOCX extraction failed: {e}")

    return TextExtract(raw_text="", extraction_method="failed")


# ══════════════════════════════════════════════════════
# 2.  Layout parsing
# ══════════════════════════════════════════════════════

# Heading heuristics: lines that look like headings
_HEADING_RE = re.compile(
    r"^(#{1,6}\s+.+|[A-Z][A-Z\s]{3,}[A-Z]|(?:\d+\.)+\s+[A-Z].{3,})$",
    re.MULTILINE,
)

def parse_layout(raw_text: str, file_bytes: bytes, filename: str) -> LayoutInfo:
    """
    Detect headings, table count, and word count from extracted text.
    For PDFs uses PyMuPDF block analysis when available.
    """
    ext = Path(filename).suffix.lower()

    headings: list[str] = []
    table_count = 0
    word_count = len(raw_text.split())

    if ext == ".pdf":
        headings, table_count = _pdf_layout(file_bytes)
    elif ext in (".docx", ".doc"):
        headings, table_count = _docx_layout(file_bytes)
    else:
        # Regex-based heading detection for plain text
        headings = _HEADING_RE.findall(raw_text)[:20]

    return LayoutInfo(
        headings=headings[:20],   # cap at 20 to keep object small
        table_count=table_count,
        word_count=word_count,
    )


def _pdf_layout(file_bytes: bytes) -> tuple[list[str], int]:
    try:
        import fitz

        doc = fitz.open(stream=file_bytes, filetype="pdf")
        headings: list[str] = []
        table_count = 0

        for page in doc:
            blocks = page.get_text("dict")["blocks"]
            for block in blocks:
                if block.get("type") != 0:
                    continue
                for line in block.get("lines", []):
                    for span in line.get("spans", []):
                        # Large or bold fonts → likely a heading
                        if span.get("size", 0) > 13 or (span.get("flags", 0) & 2**4):
                            text = span["text"].strip()
                            if text and len(text) < 200:
                                headings.append(text)

            # Simple table heuristic: pages with many aligned rects
            drawings = page.get_drawings()
            rect_count = sum(1 for d in drawings if d["type"] == "re")
            if rect_count > 8:
                table_count += 1

        doc.close()
        return headings, table_count

    except Exception as e:
        logger.warning(f"PDF layout parse failed: {e}")
        return [], 0


def _docx_layout(file_bytes: bytes) -> tuple[list[str], int]:
    try:
        import docx

        doc = docx.Document(io.BytesIO(file_bytes))
        headings = [
            p.text.strip()
            for p in doc.paragraphs
            if p.style.name.startswith("Heading") and p.text.strip()
        ]
        table_count = len(doc.tables)
        return headings, table_count

    except Exception as e:
        logger.warning(f"DOCX layout parse failed: {e}")
        return [], 0


# ══════════════════════════════════════════════════════
# 3.  Document metadata
# ══════════════════════════════════════════════════════

def extract_doc_metadata(file_bytes: bytes, filename: str) -> DocMetadata:
    """
    Extract author, creation date, modification date, and software
    from document metadata.
    """
    ext = Path(filename).suffix.lower()

    if ext == ".pdf":
        return _pdf_metadata(file_bytes)
    if ext in (".docx", ".doc"):
        return _docx_metadata(file_bytes)

    return DocMetadata()


def _pdf_metadata(file_bytes: bytes) -> DocMetadata:
    try:
        import fitz

        doc = fitz.open(stream=file_bytes, filetype="pdf")
        meta = doc.metadata
        doc.close()

        return DocMetadata(
            author           = meta.get("author") or None,
            creator_software = meta.get("creator") or meta.get("producer") or None,
            creation_date    = meta.get("creationDate") or None,
            modification_date= meta.get("modDate") or None,
            title            = meta.get("title") or None,
            subject          = meta.get("subject") or None,
        )
    except Exception as e:
        logger.warning(f"PDF metadata failed: {e}")
        return DocMetadata()


def _docx_metadata(file_bytes: bytes) -> DocMetadata:
    try:
        import docx

        doc  = docx.Document(io.BytesIO(file_bytes))
        cp   = doc.core_properties

        def safe(v) -> str | None:
            return str(v) if v else None

        return DocMetadata(
            author            = safe(cp.author),
            last_modified_by  = safe(cp.last_modified_by),
            creator_software  = safe(cp.identifier),
            creation_date     = cp.created.isoformat() if cp.created else None,
            modification_date = cp.modified.isoformat() if cp.modified else None,
            title             = safe(cp.title),
            subject           = safe(cp.subject),
            revision          = cp.revision,
        )
    except Exception as e:
        logger.warning(f"DOCX metadata failed: {e}")
        return DocMetadata()


# ══════════════════════════════════════════════════════
# Pipeline entry point
# ══════════════════════════════════════════════════════

def preprocess_document(
    file_bytes: bytes,
    source_ref: str = "uploaded_document",
) -> NormalizedFeatureObject:
    """
    Full document preprocessing pipeline.
    Returns a NormalizedFeatureObject ready for the agent committee.
    """
    errors: list[str] = []

    # ── Step 1: Text extract ──
    text_extract = extract_text(file_bytes, source_ref)
    if not text_extract.raw_text:
        errors.append(f"Text extraction failed (method: {text_extract.extraction_method})")

    # ── Step 2: Layout parse ──
    layout = parse_layout(text_extract.raw_text, file_bytes, source_ref)

    # ── Step 3: Doc metadata ──
    doc_meta = extract_doc_metadata(file_bytes, source_ref)

    # ── Dedup hash (SHA-256 of visible text) ──
    dedup_hash = hashlib.sha256(text_extract.raw_text.encode()).hexdigest()

    # ── Detect language (reuse image processor heuristic) ──
    arabic_chars = sum(1 for c in text_extract.raw_text if "\u0600" <= c <= "\u06FF")
    latin_chars  = sum(1 for c in text_extract.raw_text if c.isalpha() and ord(c) < 256)
    language = "ar" if arabic_chars > latin_chars else ("en" if latin_chars else "unknown")

    # ── Assemble DocumentData ──
    document_data = DocumentData(
        text_extract = text_extract,
        layout       = layout,
        doc_meta     = doc_meta,
        file_size_bytes = len(file_bytes),
        filename     = source_ref,
    )

    # ── Quality gate ──
    quality_passed, quality_reason = _quality_gate(text_extract, document_data)

    return NormalizedFeatureObject(
        input_type     = InputType.DOCUMENT,
        source_ref     = source_ref,
        text           = text_extract.raw_text[:10_000],   # cap for downstream
        language       = language,
        document_data  = document_data,
        quality_passed = quality_passed,
        quality_reason = quality_reason,
        dedup_hash     = dedup_hash,
        errors         = errors,
    )


def _quality_gate(
    text_extract: TextExtract,
    document_data: DocumentData,
) -> tuple[bool, str]:
    if text_extract.extraction_method == "unsupported":
        return False, "Unsupported document format"
    if text_extract.extraction_method == "failed":
        return False, "Text extraction failed"
    if len(text_extract.raw_text.strip()) < 20:
        return False, "Document has no readable content"
    return True, "OK"